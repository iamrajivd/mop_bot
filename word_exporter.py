from docx import Document

def generate_word(title, description, version, components, commands):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Version: {version}")
    doc.add_paragraph(description)
    doc.add_heading("Components and Commands", level=1)

    for comp in components:
        doc.add_heading(comp, level=2)
        for cmd in commands.get(comp, []):
            doc.add_paragraph(cmd, style='List Bullet')

    output_path = "mop_output.docx"
    doc.save(output_path)
    return output_path
